from flask import Flask, render_template, request, send_from_directory
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import datetime
import json

app = Flask(__name__)

LETTERS_FOLDER = 'generated_letters'

if not os.path.exists(LETTERS_FOLDER):
    os.makedirs(LETTERS_FOLDER)

# Load templates from a JSON file
def load_templates():
    with open('letter_templates.json', 'r') as f:
        return json.load(f)

# Generate content using the subject as the key from the loaded templates
def generate_content(subject):
    templates = load_templates()

    # Use a default template if subject is not recognized
    template = templates.get(subject, "I am writing to you regarding [Subject]. I would like to discuss [Details].")
    return template.replace('[Subject]', subject).replace('[Details]', 'the relevant details of the subject.')


def format_letter(sender_name, sender_address, recipient_name, recipient_address, subject, content):
    doc = Document()

    # Set font size and style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Add Sender details (aligned to the right)
    doc.add_paragraph(f'{sender_name}').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    doc.add_paragraph(f'{sender_address}').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add a line break
    doc.add_paragraph()

    # Add Recipient details (aligned to the left)
    doc.add_paragraph(f'{recipient_name}').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f'{recipient_address}').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add a line break
    doc.add_paragraph()

    # Add subject
    doc.add_paragraph(f'Subject: {subject}').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add a line break
    doc.add_paragraph()

    # Add Letter Content (normal alignment)
    doc.add_paragraph(content)

    # Add Signature Area
    doc.add_paragraph('\nSincerely,')
    doc.add_paragraph(f'{sender_name}')

    return doc


@app.route('/')
def index():
    return render_template('form.html')

@app.route('/generate_letter', methods=['POST'])
def generate_letter():
    # Get form data
    sender_name = request.form['sender_name']
    sender_address = request.form['sender_address']
    recipient_name = request.form['recipient_name']
    recipient_address = request.form['recipient_address']
    subject = request.form['subject']
    date = request.form['date']  # We assume date input is included

    # Generate Content based on subject
    content = generate_content(subject)

    # Generate the letter document
    doc = format_letter(sender_name, sender_address, recipient_name, recipient_address, subject, content)

    # Save the document
    file_name = f'letter_{sender_name.replace(" ", "_")}.docx'
    file_path = os.path.join(LETTERS_FOLDER, file_name)
    doc.save(file_path)

    # Return the file for download
    return send_from_directory(LETTERS_FOLDER, file_name, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)



